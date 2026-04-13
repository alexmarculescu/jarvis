import hashlib
import os
import re
from pathlib import Path
from typing import List, Tuple

import chromadb
import pandas as pd
from docx import Document
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer

# =========================
# CONFIG
# =========================
BASE_DIR = Path.home() / "jarvis"
DOCS_DIR = BASE_DIR / "docs"
DB_DIR = BASE_DIR / "db"
COLLECTION_NAME = "jarvis_docs_v2"

EMBED_MODEL = "all-MiniLM-L6-v2"
CHUNK_SIZE = 900
CHUNK_OVERLAP = 120
BATCH_SIZE = 3000

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".docx",
    ".xlsx",
    ".xls",
    ".csv",
    ".md",
}

IGNORED_FILENAMES = {
    ".DS_Store",
}

# =========================
# HELPERS
# =========================
def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    text = clean_text(text)
    if not text:
        return []

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= text_len:
            break
        start = max(end - overlap, start + 1)

    return chunks


def make_chunk_id(file_path: Path, chunk_index: int, chunk_text_value: str) -> str:
    digest = hashlib.sha1(
        f"{file_path.resolve()}::{chunk_index}::{chunk_text_value}".encode("utf-8")
    ).hexdigest()
    return digest


def classify_source(path: Path) -> str:
    p = str(path).lower()

    internal_markers = [
        "specific setup and support staff manuals",
        "cics - ",
        "quick start for support staff",
        "onboarding workflow new employee",
        "important cics queries",
        "template logic cics",
        "prepare for semester in cics",
        "academic_calendar_workflow",
        "student manuals",
        "teaching staff manuals",
    ]
    if any(marker in p for marker in internal_markers):
        return "internal"

    operational_markers = [
        ".docx",
        ".xlsx",
        ".xls",
        ".txt",
        ".csv",
        "onedrive_1_4-7-2026",
    ]
    if any(marker in p for marker in operational_markers):
        return "operational"

    return "reference"


# =========================
# FILE READERS
# =========================
def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            pages.append(f"[Page {i}]\n{page_text}")
    return clean_text("\n\n".join(pages))


def read_txt(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return clean_text(path.read_text(encoding=encoding))
        except Exception:
            continue
    return ""


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return clean_text("\n".join(parts))


def read_excel(path: Path) -> str:
    if path.suffix.lower() == ".xls":
        sheet_dict = pd.read_excel(path, sheet_name=None, dtype=str, engine="xlrd")
    else:
        sheet_dict = pd.read_excel(path, sheet_name=None, dtype=str, engine="openpyxl")

    parts = []

    for sheet_name, df in sheet_dict.items():
        if not str(sheet_name).strip():
            continue

        df = df.fillna("")
        parts.append(f"=== Sheet: {sheet_name} ===")

        headers = [str(col).strip() for col in df.columns]
        if headers:
            parts.append(" | ".join(headers))

        for _, row in df.iterrows():
            values = [str(v).strip() for v in row.tolist()]
            if any(values):
                parts.append(" | ".join(values))

        parts.append("")

    return clean_text("\n".join(parts))


def read_csv(path: Path) -> str:
    df = pd.read_csv(path, dtype=str).fillna("")
    parts = []

    headers = [str(col).strip() for col in df.columns]
    if headers:
        parts.append(" | ".join(headers))

    for _, row in df.iterrows():
        values = [str(v).strip() for v in row.tolist()]
        if any(values):
            parts.append(" | ".join(values))

    return clean_text("\n".join(parts))


def read_file(path: Path) -> str:
    ext = path.suffix.lower()

    if ext == ".pdf":
        return read_pdf(path)
    if ext in {".txt", ".md"}:
        return read_txt(path)
    if ext == ".docx":
        return read_docx(path)
    if ext in {".xlsx", ".xls"}:
        return read_excel(path)
    if ext == ".csv":
        return read_csv(path)

    return ""


# =========================
# INGEST
# =========================
def collect_files(root: Path) -> List[Path]:
    files = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if path.name in IGNORED_FILENAMES:
            continue
        if path.name.startswith("._"):
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        files.append(path)
    return sorted(files)


def build_chunks(files: List[Path]) -> Tuple[List[str], List[str], List[dict]]:
    all_ids = []
    all_docs = []
    all_metas = []

    for file_path in files:
        try:
            print(f"Reading: {file_path}")
            text = read_file(file_path)

            if not text.strip():
                print(f"Skipping empty: {file_path}")
                continue

            rel_path = str(file_path.relative_to(DOCS_DIR))
            source_group = classify_source(file_path)

            chunks = chunk_text(text)
            if not chunks:
                print(f"No chunks created: {file_path}")
                continue

            for idx, chunk in enumerate(chunks):
                chunk_with_header = f"[Source: {rel_path}]\n{chunk}"
                chunk_id = make_chunk_id(file_path, idx, chunk_with_header)

                all_ids.append(chunk_id)
                all_docs.append(chunk_with_header)
                all_metas.append(
                    {
                        "source": str(file_path),
                        "relative_path": rel_path,
                        "filename": file_path.name,
                        "extension": file_path.suffix.lower(),
                        "chunk": idx,
                        "source_group": source_group,
                    }
                )

        except Exception as exc:
            print(f"Skipping {file_path}: {exc}")

    return all_ids, all_docs, all_metas


def ingest():
    print("Loading embedder...")
    embedder = SentenceTransformer(EMBED_MODEL)

    print("Opening Chroma DB...")
    client = chromadb.PersistentClient(path=str(DB_DIR))
    collection = client.get_or_create_collection(name=COLLECTION_NAME)

    files = collect_files(DOCS_DIR)
    print(f"Found {len(files)} supported files.")

    ids, docs, metas = build_chunks(files)
    print(f"Built {len(ids)} chunks before dedup.")

    if not ids:
        print("No chunks to ingest.")
        return

    existing = set(collection.get(include=[], limit=100000).get("ids", []))
    new_rows = [
        (i, d, m)
        for i, d, m in zip(ids, docs, metas)
        if i not in existing
    ]

    if not new_rows:
        print("No new chunks to ingest.")
        return

    new_ids = [r[0] for r in new_rows]
    new_docs = [r[1] for r in new_rows]
    new_metas = [r[2] for r in new_rows]

    print(f"Ingesting {len(new_ids)} new chunks...")

    for start in range(0, len(new_ids), BATCH_SIZE):
        end = min(start + BATCH_SIZE, len(new_ids))
        batch_ids = new_ids[start:end]
        batch_docs = new_docs[start:end]
        batch_metas = new_metas[start:end]

        embeddings = embedder.encode(batch_docs, show_progress_bar=False).tolist()

        collection.add(
            ids=batch_ids,
            documents=batch_docs,
            metadatas=batch_metas,
            embeddings=embeddings,
        )

        print(f"Added batch {start} to {end} of {len(new_ids)}")

    print(f"Ingested {len(new_ids)} new chunks into collection: {COLLECTION_NAME}")


if __name__ == "__main__":
    ingest()
